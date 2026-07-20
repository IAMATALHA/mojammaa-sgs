from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Guide_utilisation_Mojammaa_Professeurs_Administration.docx"
LOGO = ROOT / "assets" / "logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B7791F"
RED = "9B1C1C"
GREEN = "1F6B4F"

FONT = "Calibri"
BODY_SIZE = 11
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
LIST_LEFT_DXA = 540
LIST_HANGING_DXA = 270
LIST_MARKER_DXA = 270


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, size=BODY_SIZE, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_style(style, size=BODY_SIZE, color=INK, bold=False, before=0, after=6, line=1.25):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def get_num_ids(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    result = {}
    for kind, fmt, text in (("bullet", "bullet", "•"), ("decimal", "decimal", "%1.")):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abstract))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(LIST_LEFT_DXA))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(LIST_LEFT_DXA))
        ind.set(qn("w:hanging"), str(LIST_HANGING_DXA))
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(next_abstract))
        num.append(abstract_ref)
        numbering.append(num)
        result[kind] = next_num
        result[f"{kind}_abstract"] = next_abstract
        next_abstract += 1
        next_num += 1
    return result


def new_num_id(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_num = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return next_num


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text="", style="Normal", bold_prefix=None, color=None, italic=False):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color or INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color or INK, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color or INK, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    apply_numbering(p, NUM_IDS["bullet"])
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_step(doc, text):
    global LAST_STEP_PARAGRAPH, LAST_STEP_NUM_ID
    p = doc.add_paragraph(style="List Number")
    previous_is_step = bool(len(doc.paragraphs) >= 2 and LAST_STEP_PARAGRAPH is doc.paragraphs[-2]._p)
    if previous_is_step:
        num_id = LAST_STEP_NUM_ID
    else:
        num_id = new_num_id(doc, NUM_IDS["decimal_abstract"])
    apply_numbering(p, num_id)
    run = p.add_run(text)
    set_run_font(run)
    LAST_STEP_PARAGRAPH = p._p
    LAST_STEP_NUM_ID = num_id
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, label, text, fill=CALLOUT, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10, color=label_color, bold=True)
    body_run = p.add_run(text)
    set_run_font(body_run, size=10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_two_col_table(doc, rows, header=None, widths=(2700, 6660)):
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    set_table_geometry(table, list(widths))
    row_index = 0
    if header:
        mark_header_row(table.rows[0])
        for col, value in enumerate(header):
            cell = table.cell(0, col)
            shade_cell(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=10, color=INK, bold=True)
        row_index = 1
    for label, detail in rows:
        first = table.cell(row_index, 0)
        second = table.cell(row_index, 1)
        p1 = first.paragraphs[0]
        p2 = second.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r2 = p2.add_run(detail)
        set_run_font(r1, size=10, color=INK, bold=True)
        set_run_font(r2, size=10, color=INK)
        row_index += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_three_col_table(doc, rows, header, widths=(1900, 3800, 3660)):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    set_table_geometry(table, list(widths))
    mark_header_row(table.rows[0])
    for col, value in enumerate(header):
        cell = table.cell(0, col)
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=10, color=INK, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            p = table.cell(row_index, col).paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=10, color=INK, bold=(col == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], size=11, color=INK, after=6, line=1.25)
    set_style(doc.styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.15)
    set_style(doc.styles["Heading 2"], size=13, color=BLUE, bold=True, before=14, after=7, line=1.15)
    set_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=10, after=5, line=1.15)
    set_style(doc.styles["List Bullet"], size=11, color=INK, after=4, line=1.25)
    set_style(doc.styles["List Number"], size=11, color=INK, after=4, line=1.25)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("Mojammaa Al Maarifa  •  Guide d’utilisation")
    set_run_font(hr, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("Professeurs & Administration  •  Page ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(fp)

    doc.core_properties.title = "Guide d’utilisation Mojammaa — Professeurs & Administration"
    doc.core_properties.subject = "Guide de prise en main de l’application Mojammaa Al Maarifa"
    doc.core_properties.author = "Mojammaa Al Maarifa"
    doc.core_properties.keywords = "Mojammaa, guide, professeur, administration, école"


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        shape = p.add_run().add_picture(str(LOGO), width=Inches(1.35))
        doc_pr = shape._inline.docPr
        doc_pr.set("title", "Logo Mojammaa Al Maarifa")
        doc_pr.set("descr", "Logo de l’établissement Mojammaa Al Maarifa")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("GUIDE DE PRISE EN MAIN")
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Mojammaa Al Maarifa")
    set_run_font(r, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Professeurs & Administration")
    set_run_font(r, size=16, color=DARK_BLUE)

    add_callout(
        doc,
        "OBJECTIF",
        "Présenter les parcours essentiels de l’application : connexion, suivi de la journée, appel, notes, devoirs, messages, statistiques et supervision.",
        fill=LIGHT_BLUE,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Document de référence interne")
    set_run_font(r, size=10, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Version de démonstration · 2026")
    set_run_font(r, size=10, color=MUTED)

    doc.add_page_break()


def build_content(doc):
    add_heading(doc, "1. Comprendre l’application", 1)
    add_para(doc, "Mojammaa Al Maarifa centralise les informations utiles à la vie scolaire. Chaque compte ouvre un espace adapté à son rôle. Ce guide se concentre sur l’espace Professeur et l’espace Direction / Administration.")
    add_two_col_table(
        doc,
        [
            ("Professeur", "Suit ses classes, fait l’appel, saisit les notes, publie des devoirs, échange avec les parents et partage des ressources."),
            ("Administration", "Supervise l’établissement, consulte les statistiques, gère les utilisateurs, les calendriers, les communications et les opérations du jour."),
            ("Données", "Les informations affichées dépendent du rôle et des autorisations du compte. Chaque utilisateur ne voit que son espace."),
        ],
        header=("Espace", "Utilité principale"),
    )

    add_heading(doc, "2. Première connexion", 1)
    add_step(doc, "Ouvrir l’application Mojammaa Al Maarifa.")
    add_step(doc, "Saisir l’adresse email du compte et le mot de passe, puis toucher « Se connecter ». ")
    add_step(doc, "Si nécessaire, toucher l’icône de langue pour choisir Français, العربية ou English.")
    add_step(doc, "Après la connexion, vérifier que l’espace affiché correspond au rôle : « Espace professeur » ou « Espace direction ». ")
    add_callout(doc, "MOT DE PASSE OUBLIÉ", "Saisir l’email, puis toucher « Mot de passe oublié ? ». Un email de réinitialisation est envoyé si le compte est reconnu.", fill="FFF7E6", label_color=GOLD)

    add_heading(doc, "3. Navigation commune", 1)
    add_para(doc, "La barre de navigation située en bas de l’écran permet d’accéder rapidement aux fonctions principales. Les onglets peuvent varier selon le rôle.")
    add_three_col_table(
        doc,
        [
            ("Accueil", "Résumé de la journée", "Priorités et raccourcis"),
            ("Messages", "Boîte de réception et envoi", "Un badge signale les non-lus"),
            ("Réglages", "Langue, compte, notifications", "Déconnexion sécurisée"),
        ],
        header=("Onglet", "Contenu", "À retenir"),
    )

    add_heading(doc, "4. Espace Professeur", 1)
    add_para(doc, "Le parcours professeur est organisé autour de la journée de cours : consulter l’EDT, faire l’appel, mettre à jour les apprentissages et communiquer avec les parents.")

    add_heading(doc, "4.1 Accueil : le focus du jour", 2)
    add_bullet(doc, "Consulter le cours en cours et le prochain cours.")
    add_bullet(doc, "Utiliser l’action principale « Faire l’appel » lorsque la séance est active.")
    add_bullet(doc, "Accéder rapidement à un devoir, à un message ou aux statistiques de ses classes.")
    add_bullet(doc, "Repérer les actions encore à traiter et les annonces récentes de la direction.")

    add_heading(doc, "4.2 Emploi du temps", 2)
    add_step(doc, "Ouvrir l’onglet « Horaires ».")
    add_step(doc, "Consulter les séances de la journée ou de la semaine.")
    add_step(doc, "Toucher une séance pour accéder à la classe et aux actions disponibles : appel ou devoir.")
    add_callout(doc, "À SAVOIR", "L’action de prière est disponible uniquement lorsque le cours concerné est réellement en cours dans l’emploi du temps.", fill=CALLOUT, label_color=DARK_BLUE)

    add_heading(doc, "4.3 Classes et appel", 2)
    add_para(doc, "Depuis « Classes », sélectionner une classe pour ouvrir son dossier. L’écran donne accès à la liste des élèves et aux actions pédagogiques.")
    add_heading(doc, "Faire l’appel", 3)
    add_step(doc, "Ouvrir la classe concernée, puis toucher « Faire l’appel ».")
    add_step(doc, "Les élèves sont présents par défaut. Toucher uniquement les élèves absents : ils deviennent rouges.")
    add_step(doc, "Vérifier la séance et la date, puis toucher « Sauvegarder l’appel ».")
    add_step(doc, "Relire le résumé enregistré. Les parents concernés peuvent être notifiés automatiquement.")
    add_bullet(doc, "Une absence déjà déclarée par un parent est identifiable dans l’écran d’appel.")
    add_bullet(doc, "En cas d’erreur, rouvrir la séance et corriger avant de sauvegarder à nouveau.")

    add_heading(doc, "4.4 Notes", 2)
    add_step(doc, "Ouvrir « Notes » depuis le dossier de classe.")
    add_step(doc, "Choisir la matière et le semestre concerné.")
    add_step(doc, "Saisir les notes sur 20, puis toucher « Sauver ».")
    add_bullet(doc, "Une note doit être comprise entre 0 et 20.")
    add_bullet(doc, "Pour une saisie en volume, utiliser « Importer un fichier Excel / CSV », vérifier l’aperçu, puis confirmer l’import.")
    add_bullet(doc, "Contrôler le nombre de contrôles et les élèves détectés avant l’enregistrement définitif.")

    add_heading(doc, "4.5 Devoirs", 2)
    add_step(doc, "Ouvrir « Devoirs », puis toucher le bouton « + ».")
    add_step(doc, "Renseigner le titre, la description, la classe cible et la date limite.")
    add_step(doc, "Ajouter si besoin une photo du tableau, une image ou un fichier PDF.")
    add_step(doc, "Toucher « Créer ». Le devoir apparaît ensuite dans la liste et dans l’espace parent.")
    add_bullet(doc, "Un devoir existant peut être réutilisé puis ajusté avant publication.")

    add_heading(doc, "4.6 Messages aux parents", 2)
    add_step(doc, "Ouvrir « Messages », puis toucher « Nouveau message ».")
    add_step(doc, "Choisir les classes destinataires : les parents des classes sélectionnées seront ciblés.")
    add_step(doc, "Renseigner l’objet et le message. Utiliser le mode urgent uniquement si nécessaire.")
    add_step(doc, "Vérifier les destinataires, puis toucher « Envoyer ».")
    add_bullet(doc, "Les messages urgents arrivent avec un préfixe et un badge distinctif.")
    add_bullet(doc, "Les messages envoyés restent consultables dans la messagerie.")

    add_heading(doc, "4.7 Comportement", 2)
    add_para(doc, "Depuis le dossier de classe ou l’écran d’appel, ouvrir l’action « Comportement » pour enregistrer un mérite ou un avertissement.")
    add_bullet(doc, "Choisir la catégorie et le motif.")
    add_bullet(doc, "Ajouter un commentaire lorsque le contexte est utile aux parents.")
    add_bullet(doc, "Enregistrer : le parent lié peut recevoir une notification.")

    add_heading(doc, "4.8 Ressources", 2)
    add_step(doc, "Ouvrir « Ressources » dans le dossier de classe.")
    add_step(doc, "Toucher « Nouvelle ressource », choisir un document ou une image et renseigner son titre.")
    add_step(doc, "Toucher « Publier ». La ressource devient visible dans l’espace parent concerné.")

    add_heading(doc, "4.9 Suivi de la prière", 2)
    add_para(doc, "Pendant le cours en cours, le professeur peut faire progresser le statut collectif de la classe : départ, arrivée à la prière, retour en classe.")
    add_bullet(doc, "Le suivi concerne la classe, pas les élèves individuellement.")
    add_bullet(doc, "Le professeur qui démarre le déplacement reste responsable de sa mise à jour.")

    add_heading(doc, "4.10 Routine quotidienne recommandée", 2)
    add_bullet(doc, "Consulter l’accueil et l’emploi du temps au début de la journée.")
    add_bullet(doc, "Faire l’appel au début de chaque séance et vérifier la sauvegarde.")
    add_bullet(doc, "Saisir les notes et publier les devoirs après le cours, avec la bonne classe et le bon semestre.")
    add_bullet(doc, "Lire les messages et répondre aux communications importantes.")

    add_heading(doc, "5. Espace Direction / Administration", 1)
    add_para(doc, "L’espace Direction donne une vue de supervision de l’établissement. Il permet de suivre l’activité du jour et d’intervenir sur l’organisation, les utilisateurs et les communications.")

    add_heading(doc, "5.1 Tableau de bord", 2)
    add_bullet(doc, "Voir la présence du jour, les appels réalisés et les classes qui attendent encore l’appel.")
    add_bullet(doc, "Repérer les absences à vérifier, les devoirs actifs et les messages du jour.")
    add_bullet(doc, "Ouvrir les priorités proposées par l’application pour accéder directement à l’écran concerné.")
    add_bullet(doc, "Accéder aux modules : statistiques, calendrier, utilisateurs, sortie/transport, prière et messagerie.")

    add_heading(doc, "5.2 Utilisateurs", 2)
    add_step(doc, "Ouvrir le module « Utilisateurs » depuis les actions d’administration.")
    add_step(doc, "Filtrer par catégorie : tous, administrateurs, professeurs ou parents.")
    add_step(doc, "Vérifier le rôle, l’activité récente et les informations du compte affiché.")
    add_callout(doc, "PRUDENCE", "Les emails, rôles et informations de compte sont des données sensibles. Ne les partagez pas dans une capture d’écran ou un groupe non autorisé.", fill="FFF7E6", label_color=GOLD)

    add_heading(doc, "5.3 Statistiques et supervision pédagogique", 2)
    add_para(doc, "L’onglet « Stats » propose plusieurs niveaux de lecture pour suivre l’établissement :")
    add_bullet(doc, "École : présence, notes, devoirs actifs et activité générale.")
    add_bullet(doc, "Classes : comparaison des moyennes, de la présence et de l’indice de santé de classe.")
    add_bullet(doc, "Matières : moyennes, taux de réussite, professeurs concernés et élèves à soutenir.")
    add_bullet(doc, "Suivi : classes à surveiller, tendance des absences et activité des professeurs.")

    add_heading(doc, "5.4 Appels, absences et retards", 2)
    add_step(doc, "Ouvrir la priorité « Appels » pour voir les classes qui ont ou n’ont pas encore enregistré l’appel.")
    add_step(doc, "Ouvrir « Absences » pour consulter les absences, les retards et les élèves récidivistes du mois.")
    add_step(doc, "Vérifier les absences déclarées par les parents et traiter les situations qui nécessitent une décision.")

    add_heading(doc, "5.5 Emploi du temps et calendrier", 2)
    add_bullet(doc, "« Emploi du temps » permet de consulter le planning par classe.")
    add_bullet(doc, "« Calendrier » permet d’ajouter un jour spécial : vacances, événement, examen ou cours annulés.")
    add_bullet(doc, "Les échéances de devoirs et les événements à venir sont regroupés dans l’agenda.")

    add_heading(doc, "5.6 Messagerie et annonces", 2)
    add_step(doc, "Ouvrir « Messages » et toucher le bouton de rédaction.")
    add_step(doc, "Choisir l’audience : tout le monde, parents, professeurs, une classe ou des personnes précises.")
    add_step(doc, "Saisir l’objet et le contenu, joindre si besoin une image ou un document, puis choisir la priorité.")
    add_step(doc, "Relire l’audience et envoyer l’annonce.")
    add_bullet(doc, "Les onglets « Envoyés » et « Supervision » permettent de suivre les communications et les échanges professeurs-parents.")
    add_bullet(doc, "Les accusés de lecture indiquent combien de destinataires ont lu un message.")
    add_bullet(doc, "Un rappel peut être envoyé aux destinataires qui n’ont pas encore lu un message.")

    add_heading(doc, "5.7 Devoirs et matières", 2)
    add_para(doc, "La Direction peut consulter les devoirs actifs et analyser le détail des résultats par matière depuis les écrans de suivi. Cette vue aide à identifier les classes ou matières qui nécessitent un accompagnement.")

    add_heading(doc, "5.8 Sortie scolaire et transport", 2)
    add_step(doc, "Ouvrir « Sortie » puis « Gérer la sortie et le transport scolaire ».")
    add_step(doc, "Ouvrir le créneau parent pendant trois heures au moment de la sortie.")
    add_step(doc, "Suivre la file en direct : « À appeler », « Appelés », « Prêts », puis « Sorties terminées ».")
    add_step(doc, "Avant de confirmer une remise, vérifier l’identité de la personne autorisée ou le code de retrait.")
    add_step(doc, "Fermer le créneau lorsque les nouvelles arrivées ne doivent plus être acceptées.")

    add_heading(doc, "5.9 Suivi de la prière", 2)
    add_para(doc, "L’écran « Prière » donne à la Direction une vue en direct des déplacements déclarés par les professeurs et des sessions terminées dans la journée.")
    add_bullet(doc, "Les états suivis sont : en route, à la prière, revenue en classe.")
    add_bullet(doc, "Le suivi est collectif et ne stocke pas de donnée individuelle sur les élèves.")

    add_heading(doc, "5.10 Routine quotidienne recommandée", 2)
    add_bullet(doc, "Commencer par le tableau de bord et traiter les priorités du jour.")
    add_bullet(doc, "Vérifier les appels manquants et les absences à traiter.")
    add_bullet(doc, "Consulter les messages entrants et suivre les annonces envoyées.")
    add_bullet(doc, "Contrôler l’emploi du temps, le calendrier et les événements à venir.")
    add_bullet(doc, "Avant la sortie, ouvrir Smart Pickup et suivre les étapes de remise.")

    add_heading(doc, "6. Démonstration recommandée à la directrice", 1)
    add_para(doc, "Pour présenter l’avancement de manière simple, suivre ce parcours de cinq à dix minutes :")
    add_step(doc, "Se connecter avec un compte Direction et montrer l’accueil : présence du jour, priorités, appels et messages.")
    add_step(doc, "Ouvrir les statistiques et montrer les vues École, Classes et Matières.")
    add_step(doc, "Ouvrir la messagerie, montrer la création d’une annonce et les accusés de lecture.")
    add_step(doc, "Présenter le parcours professeur : emploi du temps, appel, notes et devoirs.")
    add_step(doc, "Terminer par un module différenciant : sortie scolaire Smart Pickup ou suivi de la prière.")
    add_callout(doc, "PHRASE DE PRÉSENTATION", "« L’application centralise le suivi quotidien des professeurs et donne à la Direction une vue en direct de l’établissement. »", fill=LIGHT_BLUE, label_color=BLUE)

    add_heading(doc, "7. Bonnes pratiques et confidentialité", 1)
    add_bullet(doc, "Ne jamais partager un mot de passe et se déconnecter sur un appareil partagé.")
    add_bullet(doc, "Vérifier la classe, la séance, la date et le semestre avant d’enregistrer une donnée.")
    add_bullet(doc, "Utiliser les messages urgents uniquement pour les situations qui le justifient.")
    add_bullet(doc, "Ne pas diffuser les informations d’élèves, de parents ou d’utilisateurs hors des espaces prévus.")
    add_bullet(doc, "Pour une remise Smart Pickup, vérifier l’identité avant de confirmer l’opération : cette action est irréversible.")
    add_bullet(doc, "En cas d’incertitude, ne pas modifier la donnée : contacter l’administration ou l’administrateur de l’application.")

    doc.add_page_break()
    add_heading(doc, "8. Dépannage rapide", 1)
    add_two_col_table(
        doc,
        [
            ("Impossible de se connecter", "Vérifier l’email, le mot de passe et la connexion Internet. Utiliser « Mot de passe oublié ? » si nécessaire."),
            ("Aucune classe ou aucun cours", "Le profil ou l’emploi du temps doit être complété par l’administration."),
            ("Les données ne se chargent pas", "Vérifier la connexion puis utiliser « Réessayer ». Éviter d’enregistrer plusieurs fois la même action."),
            ("Un parent n’est pas trouvé", "Vérifier que le compte parent est lié à l’élève et que la classe sélectionnée est correcte."),
            ("Un fichier ne s’importe pas", "Vérifier le format Excel/CSV, les colonnes attendues et les notes comprises entre 0 et 20."),
            ("Une erreur persiste", "Noter l’écran, l’action effectuée et l’heure, puis transmettre ces éléments à l’administrateur."),
        ],
        header=("Situation", "Réponse recommandée"),
    )

    add_heading(doc, "9. Résumé en une page", 1)
    add_three_col_table(
        doc,
        [
            ("Professeur", "EDT → Appel → Notes → Devoirs", "Mettre à jour les données au plus près du cours"),
            ("Professeur", "Messages → Comportement → Ressources", "Communiquer avec précision et garder une trace"),
            ("Direction", "Accueil → Appels → Absences", "Commencer par les priorités du jour"),
            ("Direction", "Stats → Utilisateurs → Calendrier", "Superviser l’établissement et son organisation"),
            ("Direction", "Messages → Smart Pickup → Prière", "Coordonner les opérations en direct"),
        ],
        header=("Rôle", "Parcours", "Principe"),
    )
    add_callout(doc, "FIN DU GUIDE", "Ce document peut être complété avec des captures d’écran après la démonstration de la version présentée.", fill=LIGHT_GRAY, label_color=MUTED)


def main():
    global NUM_IDS, LAST_STEP_PARAGRAPH, LAST_STEP_NUM_ID
    doc = Document()
    setup_document(doc)
    NUM_IDS = get_num_ids(doc)
    LAST_STEP_PARAGRAPH = None
    LAST_STEP_NUM_ID = None
    add_cover(doc)
    build_content(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
